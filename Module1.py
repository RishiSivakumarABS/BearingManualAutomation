import streamlit as st
from docx import Document
from datetime import datetime
import pandas as pd
import io

st.set_page_config(page_title="ABS Bearing Design Tool", layout="wide")
st.title("🛠️ ABS Bearing Design Automation Tool")
st.markdown("This tool assists in selecting bearing specifications and calculating tolerance deviations based on ABS internal standards.")

# ----------------------------
# Module 3: Material & Heat Treatment Selector
# ----------------------------
def suggest_material_and_treatment_module3(
    roller_dia,
    wall_thickness,
    load_type="standard",
    ring_position=None,
    bearing_type=None,
    mill_type=None  # Add this if you want to pass it too
):
    load_type = load_type.lower()

    # 1️⃣ Priority: Load type override
    if load_type == "impact":
        return "G20Cr2Ni4A", "Carburizing Heat Treatment"

    # 2️⃣ Priority: Mill type override
    if mill_type == "hot mill":
        return "GCr18Mo", "Bainite Isothermal QT"

    # 3️⃣ Then normal stress factor logic
    if ring_position == "Inner Ring":
        if bearing_type == "Fixed":
            if roller_dia > 45 or wall_thickness > 25:
                return "GCr18Mo", "Bainite Isothermal QT"
            else:
                return "GCr15SiMn", "Martensitic Quenching"
        elif bearing_type == "Floating":
            return "GCr15", "Martensitic Quenching"

    elif ring_position == "Outer Ring":
        if bearing_type == "Fixed":
            return "GCr18Mo", "Bainite Isothermal QT"
        else:
            return "GCr15SiMn", "Martensitic Quenching"

    # 4️⃣ Fallback (legacy)
    if wall_thickness <= 17 and roller_dia <= 32:
        return "GCr15", "Martensitic Quenching"
    elif wall_thickness <= 25 and roller_dia <= 45:
        return "GCr15", "Bainite Isothermal QT"
    elif roller_dia > 45:
        return "GCr18Mo", "Bainite Isothermal QT"
    else:
        return "GCr15SiMn", "Martensitic Quenching"



# ----------------------------
# Module 4: Roller Profile Table
# ----------------------------
roller_profile_df = pd.DataFrame([
    {"Profile Type": "Logarithmic", "Min Dia (mm)": 20, "Max Dia (mm)": 40, "Max Deviation (µm)": 3.0},
    {"Profile Type": "Logarithmic", "Min Dia (mm)": 40, "Max Dia (mm)": 60, "Max Deviation (µm)": 4.0},
    {"Profile Type": "Crowned",     "Min Dia (mm)": 20, "Max Dia (mm)": 40, "Max Deviation (µm)": 2.0},
    {"Profile Type": "Crowned",     "Min Dia (mm)": 40, "Max Dia (mm)": 60, "Max Deviation (µm)": 3.0},
    {"Profile Type": "Flat",        "Min Dia (mm)": 20, "Max Dia (mm)": 60, "Max Deviation (µm)": 1.0},
])

def get_max_deviation(profile_type, diameter):
    df = roller_profile_df
    for _, row in df.iterrows():
        if row['Profile Type'].lower() == profile_type.lower() and row['Min Dia (mm)'] <= diameter <= row['Max Dia (mm)']:
            return row['Max Deviation (µm)']
    return None

# ----------------------------
# Module 2: Load Tolerance Tables
# ----------------------------
@st.cache_data
def load_tolerance_tables():
    df_normal = pd.read_excel("Table1_Normal_Tolerances.xlsx")
    df_p6 = pd.read_excel("Table2_P6_Tolerances.xlsx")
    df_p5 = pd.read_excel("Table3_P5_Tolerances.xlsx")
    return df_normal, df_p6, df_p5

df_normal, df_p6, df_p5 = load_tolerance_tables()
class_tables = {"Normal": df_normal, "P6": df_p6, "P5": df_p5}

def find_tolerance(bore_dia, tolerance_class):
    df = class_tables.get(tolerance_class)
    if df is None or df.empty:
        return None

    for _, row in df.iterrows():
        if row['Min Diameter (mm)'] < bore_dia <= row['Max Diameter (mm)']:
            upper_dev = row['Upper Deviation (µm)']
            lower_dev = row['Lower Deviation (µm)']
            max_bore_mm = bore_dia + (upper_dev / 1000)  # µm to mm
            min_bore_mm = bore_dia + (lower_dev / 1000)

            return {
                "upper_dev": upper_dev,
                "lower_dev": lower_dev,
                "max_bore": round(max_bore_mm, 3),
                "min_bore": round(min_bore_mm, 3)
            }

    return None


# ----------------------------
# Module 5: Clearance Class Checker
# ----------------------------
def suggest_clearance(bore_dia, mill_type=None):
    # Mill type overrides
    if mill_type == "hot mill":
        return "C4"
    elif mill_type == "cold mill":
        return "C3"
    
    # Default logic based on bore diameter
    if bore_dia <= 120:
        return "C2 or Normal"
    elif bore_dia <= 250:
        return "Normal or C3"
    elif bore_dia <= 500:
        return "C3 or C4"
    else:
        return "C4 or C5"


# ----------------------------
# Tabs for Modules 1 to 6
# ----------------------------
tabs = st.tabs([
    "Smart Specification Selector", 
    "Tolerance & Fit Calculator",
    "Roller Profile Matching",
    "Material & Heat Treatment Selector",
    "Clearance Class Checker",
    "Final Compliance Validator"
])

# ----------------------------
# Module 1 – Smart Specification Selector
# ----------------------------
with tabs[0]:
    st.header("🔧 Module 1: Smart Specification Selector")
    bore = st.number_input("Bore Diameter (mm)", value=250, key="mod1_bore")
    wall = st.number_input("Wall Thickness (mm)", value=20, key="mod1_wall")
    roller = st.number_input("Roller Diameter (mm)", value=35, key="mod1_roller")
    app = st.selectbox("Application Type", ["standard", "precision", "high load"], key="mod1_app")
    rpm = st.number_input("Operating Speed (RPM)", value=300, key="mod1_rpm")
    mill = st.selectbox("Mill Type (optional)", [None, "hot mill", "cold mill"], key="mod1_mill")
    load = st.selectbox("Load Type", ["standard", "impact"], key="mod1_load")
    ring_position = st.selectbox("Ring Position", ["Inner Ring", "Outer Ring"], key="mod1_ringpos")
    bearing_type = st.selectbox("Bearing Type", ["Fixed", "Floating"], key="mod1_type")


    def bearing_class(app_type): return "P5" if app_type == "precision" else "P6"

    def cage(app_type, rpm_val):
        if app_type == "high load": return "Pin-Type", "Steel"
        elif app_type == "standard" and rpm_val > 1000: return "Polymer", "Nylon/PTFE"
        elif app_type == "standard": return "Riveted", "Steel, Mass"
        else: return "Machined", "Steel, Mass"

    if st.button("Generate Specification Recommendation"):
        bc = bearing_class(app)
        cc = suggest_clearance(bore, mill)
        steel, ht = suggest_material_and_treatment_module3(
    roller, wall, load, ring_position=ring_position, bearing_type=bearing_type, mill_type=mill
)

        ct, cm = cage(app, rpm)

        st.write(f"**Bearing Class:** {bc}")
        st.write(f"**Clearance Class:** {cc}")
        st.write(f"**Steel Type:** {steel}")
        st.write(f"**Heat Treatment:** {ht}")
        st.write(f"**Cage Type & Material:** {ct} ({cm})")
        st.success("✅ Recommendation generated.")

        # Optional: generate report only for Module 1
        doc = Document()
        doc.add_heading('ABS Bearing Design Report', level=1)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        doc.add_heading('Module 1 – Specification Output', level=2)
        doc.add_paragraph(f"Bore Diameter: {bore} mm")
        doc.add_paragraph(f"Wall Thickness: {wall} mm")
        doc.add_paragraph(f"Roller Diameter: {roller} mm")
        doc.add_paragraph(f"Application Type: {app}")
        doc.add_paragraph(f"Operating Speed: {rpm} RPM")
        doc.add_paragraph(f"Mill Type: {mill}")
        doc.add_paragraph(f"Load Type: {load}")

        doc.add_paragraph(f"Bearing Class: {bc}")
        doc.add_paragraph(f"Clearance Class: {cc}")
        doc.add_paragraph(f"Steel Type: {steel}")
        doc.add_paragraph(f"Heat Treatment: {ht}")
        doc.add_paragraph(f"Cage Type & Material: {ct} ({cm})")
        doc.add_paragraph(f"Ring Position: {ring_position}")
        doc.add_paragraph(f"Bearing Type: {bearing_type}")



        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button("📥 Download Module 1 Report", data=buffer, file_name="Bearing_Module1_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ----------------------------
# Module 2 – Tolerance & Fit
# ----------------------------
with tabs[1]:
    st.header("📏 Module 2: Tolerance & Fit Calculator")

    dia2 = st.number_input("Enter Bore Diameter (mm)", value=280.0, key="mod2_dia")
    tol_class = st.selectbox("Tolerance Class", ["Normal", "P6", "P5"], key="mod2_class")

    if st.button("Calculate Tolerances", key="btn_mod2"):
        result = find_tolerance(dia2, tol_class)

        if result:
            st.success("✅ Tolerance Found:")
            st.write(f"**Upper Deviation:** +{result['upper_dev']} µm")
            st.write(f"**Lower Deviation:** {result['lower_dev']} µm")
            st.write(f"**Maximum Bore Diameter:** {result['max_bore']} mm")
            st.write(f"**Minimum Bore Diameter:** {result['min_bore']} mm")
            st.caption("Tolerance values based on ISO 492 standards")
        else:
            st.error("⚠️ Bore diameter not found in the selected tolerance class table.")


# ----------------------------
# Module 3 – Roller Profile Matching
# ----------------------------
with tabs[2]:
    st.header("📊 Module 3: Roller Profile Matching")
    
    ptype = st.selectbox("Roller Profile Type", ["Logarithmic", "Crowned", "Flat"], key="mod3_type")
    pr_dia = st.number_input("Roller Diameter (mm)", value=40.0, key="mod3_dia")
    measured_dev = st.number_input("Actual Measured Profile Deviation (µm)", value=0.0, key="mod3_meas_dev")

    if st.button("Check Roller Profile Tolerance", key="btn_mod3"):
        max_dev = get_max_deviation(ptype, pr_dia)

        if max_dev is not None:
            st.write(f"📌 Allowed Max Deviation for this configuration: **{max_dev} µm**")
            st.write(f"📌 Your Measured Deviation: **{measured_dev} µm**")

            if measured_dev <= max_dev:
                st.success("✅ PASS: Measured deviation is within the allowed tolerance.")
            else:
                st.error("❌ FAIL: Measured deviation exceeds the allowed tolerance.")
        else:
            st.error("⚠️ No tolerance data found for the selected profile and diameter.")

# ----------------------------
# Module 4 – Material Selector
# ----------------------------
with tabs[3]:
    st.header("⚙️ Module 4: Material & Heat Treatment Selector")

    rd4 = st.number_input("Roller Diameter (mm)", value=35.0, key="mod4_roller")
    wt4 = st.number_input("Wall Thickness (mm)", value=20.0, key="mod4_wall")
    load4 = st.selectbox("Load Type", ["standard", "impact"], key="mod4_load")
    ring4 = st.selectbox("Ring Position", ["Inner Ring", "Outer Ring"], key="mod4_ring")
    type4 = st.selectbox("Bearing Type", ["Fixed", "Floating"], key="mod4_type")
    mill4 = st.selectbox("Mill Type (optional)", [None, "hot mill", "cold mill"], key="mod4_mill")

    if st.button("Get Recommendation", key="btn_mod4"):
        steel4, ht4 = suggest_material_and_treatment_module3(
            rd4, wt4, load4, ring_position=ring4, bearing_type=type4, mill_type=mill4
        )
        st.success("✅ Material Recommendation Found:")
        st.write(f"**Steel Type:** {steel4}")
        st.write(f"**Heat Treatment:** {ht4}")

# ----------------------------
# Module 5 – Clearance Checker
# ----------------------------
with tabs[4]:
    st.header("📐 Module 5: Clearance Class Checker")
    
    bd5 = st.number_input("Bore Diameter (mm)", value=250.0, key="mod5_bore")
    mt5 = st.selectbox("Mill Type (optional)", [None, "hot mill", "cold mill"], key="mod5_mill")

    if st.button("Check Clearance", key="btn_mod5"):
        cc5 = suggest_clearance(bd5, mt5)
        st.success(f"✅ Recommended Clearance Class: {cc5}")


# ----------------------------
# Module 6 – Final Compliance Validator
# ----------------------------
with tabs[5]:
    st.header("✅ Module 6: Final Compliance Validator")
    st.markdown("Enter selected parameters from Modules 1–5 to validate against ABS standards.")

    f_bearing_class = st.selectbox("Selected Bearing Class", ["P5", "P6"], key="mod6_bc")
    f_clearance_class = st.selectbox("Selected Clearance Class", ["C2", "Normal", "C3", "C4", "C5"], key="mod6_cc")
    f_tol_class = st.selectbox("Selected Tolerance Class", ["Normal", "P6", "P5"], key="mod6_tol")
    f_steel = st.selectbox("Selected Steel", ["GCr15", "GCr15SiMn", "GCr18Mo", "G20Cr2Ni4A"], key="mod6_steel")
    f_ht = st.selectbox("Selected Heat Treatment", [
        "Martensitic Quenching", "Bainite Isothermal QT", "Carburizing Heat Treatment"
    ], key="mod6_ht")
    f_cage = st.selectbox("Selected Cage Type", ["Pin-Type", "Polymer", "Riveted", "Machined"], key="mod6_cage")

    if st.button("Run Compliance Check", key="mod6_check"):
        issues = []

        if f_bearing_class == "P5" and f_tol_class == "Normal":
            issues.append("P5 bearing class typically should not use Normal tolerance.")
        if f_cage == "Polymer" and f_ht == "Carburizing Heat Treatment":
            issues.append("Polymer cages are not ideal for carburized components.")

        if issues:
            st.error("❌ Compliance Issues Found:")
            for i in issues:
                st.write(f"- {i}")
        else:
            st.success("✅ All selections are compliant with defined engineering rules.")
